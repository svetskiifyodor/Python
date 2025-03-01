from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Открываем существующий документ
doc = Document("document_with_table.docx")

# Добавляем изображение
image_path = "image.jpg"  # Укажите путь к изображению
doc.add_picture(image_path, width=Inches(2))  # Устанавливаем ширину изображения (4 дюйма)

# Добавляем подпись под изображением
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание подписи по центру
para.add_run("Подпись к изображению: ОПИСАНИЕ.")  # Текст подписи

# Сохраняем изменения в новый файл
doc.save("document_with_table_and_image.docx")
