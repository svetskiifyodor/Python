from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# Создание документа
doc = Document()

# Функция для установки интервала и отступов только для маркированного списка
def set_paragraph_spacing(paragraph, indent=False):
    paragraph.paragraph_format.space_before = 0  # Убираем интервал перед абзацем
    paragraph.paragraph_format.space_after = 0   # Убираем интервал после абзаца
    paragraph.paragraph_format.line_spacing = 1  # Устанавливаем одинарный интервал
    if indent:
        paragraph.paragraph_format.left_indent = Pt(56.7)  # Устанавливаем отступ слева 2 см

# Добавление начала текста
para = doc.add_paragraph("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:")
set_paragraph_spacing(para)

# Добавление пустого абзаца (отступ после текста)
doc.add_paragraph()

# Добавление маркированного списка с прозрачными точками (имитация)
para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Флеш-память: используется для хранения скетчей.")
set_paragraph_spacing(para, indent=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("ОЗУ (")
run.bold = False  # Жирный шрифт для "SRAM"
para.add_run("SRAM").bold = True
run = para.add_run(" - ")
run.italic = True  # Курсив для фразы
para.add_run("static random access memory").italic = True
run = para.add_run("): используется для хранения и работы переменных.")
set_paragraph_spacing(para, indent=True)

para = doc.add_paragraph(style='List Bullet')
para.add_run("EEPROM (энергонезависимая память): используется для хранения постоянной информации.")
set_paragraph_spacing(para, indent=True)

# Добавление пустого абзаца после перечисления
doc.add_paragraph()

para = doc.add_paragraph(
    "Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью."
)
set_paragraph_spacing(para)

# Добавление пустого абзаца после текста
doc.add_paragraph()

# Добавление таблицы
table = doc.add_table(rows=4, cols=5)  # Уменьшаем количество строк, чтобы избежать пустой строки в конце

# Установка стиля таблицы
table.style = 'Table Grid'

# Заполнение заголовков столбцов с жирным шрифтом и серой заливкой
hdr_cells = table.rows[0].cells
hdr_cells[1].text = "ATmega168"
hdr_cells[2].text = "ATmega328"
hdr_cells[3].text = "ATmega1280"
hdr_cells[4].text = "ATmega2560"

# Устанавливаем жирный шрифт и серую заливку для названий столбцов
for cell in hdr_cells[1:]:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
    # Заливка ячеек серым цветом
    cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
    cell._element.xpath('.//w:shd')[0].set(qn('w:fill'), 'D3D3D3')  # Серый цвет
    # Выравнивание по центру
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Заполнение строк с данными
data = [
    ("Flash\n(1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"),
    ("SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"),
    ("EEPROM", "512 байт", "1024 байт", "4 Кбайт", "4 Кбайт"),
]

for i, row in enumerate(data, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
    row_cells[3].text = row[3]
    row_cells[4].text = row[4]

    # Установим жирный текст для первого столбца (названия строк) и серую заливку
    for j, cell in enumerate(row_cells):
        if j == 0:
            # Устанавливаем жирный шрифт для первого столбца
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            # Заливка ячеек серым цветом
            cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
            cell._element.xpath('.//w:shd')[0].set(qn('w:fill'), 'D3D3D3')  # Серый цвет
        # Выравнивание текста по центру
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Заливка для первой пустой ячейки (верхний левый угол таблицы)
first_cell = table.cell(0, 0)  # Первая ячейка
first_cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
first_cell._element.xpath('.//w:shd')[0].set(qn('w:fill'), 'D3D3D3')  # Серый цвет

# Добавление пустой строки после таблицы
doc.add_paragraph()

# Добавление текста после таблицы с курсивом
para = doc.add_paragraph(
    "Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. "
    "Эти данные не распространяются на операции чтения данных из EEPROM - чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.",
    style='Normal'
)

# Делаем весь абзац курсивным
for run in para.runs:
    run.italic = True

# Сохранение документа
doc.save("document_with_table.docx")
